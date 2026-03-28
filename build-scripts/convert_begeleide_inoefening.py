#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convert begeleide inoefening vragen.docx + antwoorden.docx to interactive HTML.
Usage: python convert_begeleide_inoefening.py --all

HOW TO ADAPT
============
Reads both the vragen and antwoorden Word docs and generates a single interactive
HTML page with collapsible sidebar per opgave, toggleable hints/denkstappen/formules,
and hidden answer reveals.
"""
import sys, io, os, glob, html as html_mod, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

DOMAIN_MAP = {
    'Wiskundig':        ('wiskunde',   '#1A5276', '#EBF5FB', '#154360'),
    'Economisch':       ('economisch', '#E67E22', '#FEF5E7', '#BA6A1C'),
    'Grafisch':         ('grafisch',   '#1E8449', '#E8F8F0', '#186A3B'),
    'Bedrijfseconomie': ('bedrijf',    '#E67E22', '#FEF5E7', '#BA6A1C'),
    'Marktanalyse':     ('markt',      '#1A5276', '#EBF5FB', '#154360'),
    'Marktwerking':     ('marktwerk',  '#1A5276', '#EBF5FB', '#154360'),
    'Overheidsbeleid':  ('overheid',   '#1E8449', '#E8F8F0', '#186A3B'),
    'Arbeidsmarkt':     ('arbeids',    '#E67E22', '#FEF5E7', '#BA6A1C'),
    'Externe effecten': ('extern',     '#1E8449', '#E8F8F0', '#186A3B'),
}

def get_domain_info(domain_name):
    if domain_name in DOMAIN_MAP:
        return DOMAIN_MAP[domain_name]
    slug = domain_name.lower().replace(' ', '')[:8]
    return (slug, '#1A5276', '#EBF5FB', '#154360')

def esc(text):
    return html_mod.escape(text)

def make_id(text):
    """Create a URL-safe id from question label."""
    # Extract question number like "1a", "2b", "3g-h"
    m = re.search(r'(\d+[a-z](?:-[a-z])?)', text.lower().replace(' ', ''))
    if m:
        return 'q' + m.group(1).replace('-', '')
    # fallback
    return 'q' + re.sub(r'[^a-z0-9]', '', text.lower())[:15]


def long_path(p):
    """Add \\\\?\\ prefix for Windows long paths (>260 chars)."""
    absp = os.path.abspath(p)
    if len(absp) > 250 and os.name == 'nt':
        return chr(92)*2 + '?' + chr(92) + absp
    return p

def parse_bi_document(docx_path, is_answers=False):
    """Parse a begeleide inoefening document into structured data."""
    try:
        doc = Document(long_path(docx_path))
    except Exception:
        return None

    body = doc.element.body
    stream = []
    table_index = 0
    para_index = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            if para_index < len(doc.paragraphs):
                p = doc.paragraphs[para_index]
                para_index += 1
                text = p.text.strip()
                if not text:
                    continue
                style = p.style.name if p.style else ''
                if style == 'Heading 1':
                    stream.append(('heading1', text))
                elif style == 'Heading 2':
                    stream.append(('heading2', text))
                elif style == 'List Paragraph':
                    stream.append(('list_item', text))
                else:
                    stream.append(('paragraph', text))
        elif tag == 'tbl':
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                table_index += 1
                nrows = len(table.rows)
                ncols = len(table.columns)

                # Opgave header (1×2, first cell is digit)
                if nrows == 1 and ncols == 2:
                    first = table.rows[0].cells[0].text.strip()
                    second = table.rows[0].cells[1].text.strip()
                    if first.isdigit() and '\n' in second:
                        parts = second.split('\n')
                        stream.append(('opgave_header', {
                            'nr': first,
                            'title': parts[0].strip(),
                            'domain': parts[-1].strip() if len(parts) > 1 else 'Marktanalyse',
                        }))
                        continue

                # Domain badges (1×3, skip)
                if nrows == 1 and ncols == 3:
                    continue

                # Samenvatting table
                if nrows > 1 and ncols == 2:
                    h0 = table.rows[0].cells[0].text.strip()
                    if '■' in h0 or 'Samenvatting' in h0 or '📋' in h0:
                        rows = [(r.cells[0].text.strip(), r.cells[1].text.strip()) for r in table.rows[1:]]
                        stream.append(('samenvatting', rows))
                        continue

                # Single-cell tables: helpers, formulas, callouts, answers
                if nrows == 1 and ncols == 1:
                    text = table.rows[0].cells[0].text.strip()
                    if text.startswith('▶ Denkstappen:') or text.startswith('▶ Denkstappen\n'):
                        stream.append(('helper_denkstappen', text))
                    elif text.startswith('▶ Hint:') or text.startswith('▶ Hint\n'):
                        stream.append(('helper_hint', text))
                    elif text.startswith('▶ Formule-herinnering:') or text.startswith('▶ Formule-herinnering\n') or text.startswith('▶ Formule'):
                        stream.append(('helper_formule', text))
                    elif text.startswith('▶ Werkwijze:') or text.startswith('▶ Werkwijze\n'):
                        stream.append(('helper_hint', text))  # treat as hint
                    elif text.startswith('▶ Studietip:') or text.startswith('▶ Studietip\n'):
                        continue  # skip meta instruction
                    elif text.startswith('▶ Uitleg:') or text.startswith('▶ Uitleg\n'):
                        stream.append(('uitleg', text))
                    elif text.startswith('⚠'):
                        stream.append(('helper_letop', text))
                    elif text.startswith('✅'):
                        stream.append(('helper_controle', text))
                    elif text.startswith('⭐'):
                        stream.append(('helper_kernregel', text))
                    elif text.startswith('✨'):
                        stream.append(('helper_tip', text))
                    elif is_answers and not text.startswith('▶'):
                        # In answers doc: standalone content = answer
                        stream.append(('answer', text))
                    else:
                        stream.append(('context_box', text))
                    continue

                # Multi-row tables: could be context/formulas
                if nrows > 1 or ncols > 1:
                    # Skip inhoudsopgave-like tables
                    if ncols >= 4 and table.rows[0].cells[0].text.strip() in ('Nr.', 'Nr'):
                        continue
                    stream.append(('context_box', '\n'.join(
                        ' | '.join(c.text.strip() for c in r.cells)
                        for r in table.rows
                    )))

    return stream


def merge_vragen_antwoorden(vragen_stream, antwoorden_stream):
    """Merge vragen and antwoorden streams into opgaven structure."""

    # First pass: extract opgaven and questions from vragen
    opgaven = []
    current_opgave = None
    current_question = None
    samenvatting = None
    pre_opgave_context = []  # context before first opgave

    for etype, edata in vragen_stream:
        if etype == 'opgave_header':
            if current_question:
                if current_opgave:
                    current_opgave['questions'].append(current_question)
                current_question = None
            if current_opgave:
                opgaven.append(current_opgave)
            current_opgave = {
                'nr': edata['nr'],
                'title': edata['title'],
                'domain': edata['domain'],
                'context': [],
                'questions': [],
                'intro': [],
            }
            continue

        if etype == 'heading2' and ('Vraag' in edata or 'vraag' in edata):
            if current_question and current_opgave:
                current_opgave['questions'].append(current_question)
            # Parse question label
            label = edata
            current_question = {
                'label': label,
                'id': make_id(label),
                'text': '',
                'fill_in': [],
                'helpers': [],
                'answer': None,
                'uitleg': None,
            }
            continue

        if etype == 'samenvatting':
            samenvatting = edata
            continue

        if etype == 'heading1' and 'Samenvattend' in edata:
            continue  # skip the heading, samenvatting table follows

        if current_question:
            if etype == 'paragraph':
                if not current_question['text']:
                    current_question['text'] = edata
                else:
                    current_question['fill_in'].append(edata)
            elif etype in ('helper_denkstappen', 'helper_hint', 'helper_formule',
                          'helper_letop', 'helper_controle', 'helper_kernregel', 'helper_tip'):
                current_question['helpers'].append((etype, edata))
            elif etype == 'context_box':
                # Could be fill-in template or formula reminder
                current_question['fill_in'].append(edata)
        elif current_opgave:
            if etype == 'context_box':
                current_opgave['context'].append(edata)
            elif etype == 'paragraph':
                current_opgave['intro'].append(edata)
            elif etype in ('helper_denkstappen', 'helper_hint', 'helper_formule',
                          'helper_letop', 'helper_controle'):
                # Opgave-level helper, skip (usually the how-to section)
                pass

    if current_question and current_opgave:
        current_opgave['questions'].append(current_question)
    if current_opgave:
        opgaven.append(current_opgave)

    # Second pass: extract answers from antwoorden and match to questions
    if antwoorden_stream:
        current_q_idx = 0
        current_opg_idx = 0
        ans_questions = []  # flat list of (label, answer, uitleg, helpers)
        current_ans_q = None

        for etype, edata in antwoorden_stream:
            if etype == 'opgave_header':
                continue
            if etype == 'heading2' and ('Vraag' in edata or 'vraag' in edata):
                if current_ans_q:
                    ans_questions.append(current_ans_q)
                current_ans_q = {
                    'label': edata,
                    'id': make_id(edata),
                    'answer': None,
                    'uitleg': None,
                    'helpers': [],
                }
                continue

            if current_ans_q:
                if etype == 'answer':
                    current_ans_q['answer'] = edata
                elif etype == 'uitleg':
                    current_ans_q['uitleg'] = edata
                elif etype in ('helper_denkstappen', 'helper_hint', 'helper_formule',
                              'helper_letop', 'helper_controle', 'helper_kernregel', 'helper_tip'):
                    current_ans_q['helpers'].append((etype, edata))
                elif etype == 'context_box':
                    # In answers, standalone context boxes are likely answers
                    if not current_ans_q['answer']:
                        current_ans_q['answer'] = edata

        if current_ans_q:
            ans_questions.append(current_ans_q)

        # Match answers to questions by id
        ans_by_id = {aq['id']: aq for aq in ans_questions}
        for opg in opgaven:
            for q in opg['questions']:
                if q['id'] in ans_by_id:
                    aq = ans_by_id[q['id']]
                    q['answer'] = aq.get('answer')
                    q['uitleg'] = aq.get('uitleg')
                    # Merge helpers (answers doc may have additional ones)
                    existing_types = {h[0] for h in q['helpers']}
                    for h in aq.get('helpers', []):
                        if h[0] not in existing_types:
                            q['helpers'].append(h)

    return opgaven, samenvatting


def render_helper_content(text):
    """Strip prefix and render helper text as HTML."""
    # Remove ▶ prefix and label
    clean = text
    for prefix in ['▶ Denkstappen:\n', '▶ Denkstappen:', '▶ Hint:', '▶ Hint:\n',
                    '▶ Formule-herinnering:\n', '▶ Formule-herinnering:',
                    '▶ Werkwijze:\n', '▶ Werkwijze:', '▶ Formule:\n', '▶ Formule:',
                    '▶ Uitleg:\n', '▶ Uitleg:', '▶ ']:
        if clean.startswith(prefix):
            clean = clean[len(prefix):].strip()
            break
    return clean.replace('\n', '<br>')


def render_callout_text(text):
    """Strip emoji and render callout."""
    clean = text
    for prefix in ['⚠ ', '✅ ', '⭐ ', '✨ ', '⚠', '✅', '⭐', '✨']:
        if clean.startswith(prefix):
            clean = clean[len(prefix):].strip()
            break
    return clean


def generate_html(opgaven, samenvatting, para_number, para_name):
    """Generate the full HTML."""

    # Determine primary domain color
    primary_domain = 'Marktanalyse'
    if opgaven:
        primary_domain = opgaven[0].get('domain', 'Marktanalyse')
    di = get_domain_info(primary_domain)
    dcls, dprimary, dlight, ddark = di

    # Build sidebar
    sidebar_items = ''
    for opg in opgaven:
        q_links = ''
        for q in opg['questions']:
            # Short label for sidebar
            short = q['label'].replace('Vraag ', '').replace(' — ', ' · ')
            if len(short) > 40:
                short = short[:37] + '...'
            q_links += f'        <a class="nav-q" href="#{q["id"]}" data-q="{q["id"]}">{esc(short)}</a>\n'

        expanded = ' expanded' if opg == opgaven[0] else ''
        sidebar_items += f'''    <div class="nav-opgave{expanded}">
      <div class="nav-opgave-title" data-toggle="opgave" data-scroll="opgave-{opg['nr']}">
        <span class="nav-dot"></span>
        <span class="nav-opgave-label">{opg['nr']} &middot; {esc(opg['title'][:30])}</span>
        <svg class="nav-arrow" viewBox="0 0 14 14"><polyline points="3 5 7 9 11 5"/></svg>
      </div>
      <div class="nav-questions">
{q_links}      </div>
    </div>
'''

    # Hero cards
    hero_cards = ''
    for opg in opgaven:
        nq = len(opg['questions'])
        hero_cards += f'''          <div class="hero-card" data-target="opgave-{opg['nr']}">
            <div class="hero-card-num">Opgave {opg['nr']}</div>
            <div class="hero-card-title">{esc(opg['title'][:35])}</div>
            <div class="hero-card-count">{nq} deelvra{"ag" if nq == 1 else "gen"}</div>
          </div>
'''

    # Opgave sections
    opgave_sections = ''
    for opg in opgaven:
        # Context boxes
        ctx = ''
        for c in opg.get('context', []):
            lines = c.replace('\n', '<br>')
            ctx += f'        <div class="context-box">{lines}</div>\n'

        # Intro
        intro = ''
        for t in opg.get('intro', []):
            intro += f'        <p class="opgave-intro">{esc(t)}</p>\n'

        # Questions
        questions_html = ''
        for q in opg['questions']:
            # Fill-in
            fill_in = ''
            if q['fill_in']:
                lines = '<br>'.join(esc(f) for f in q['fill_in'])
                fill_in = f'          <div class="fill-in">{lines}</div>\n'

            # Helpers
            helpers = ''
            if q['helpers']:
                helper_items = ''
                for htype, htext in q['helpers']:
                    css_map = {
                        'helper_denkstappen': ('helper-denkstappen', 'Denkstappen'),
                        'helper_hint': ('helper-hint', 'Hint'),
                        'helper_formule': ('helper-formule', 'Formule-herinnering'),
                        'helper_letop': ('helper-letop', 'Let op'),
                        'helper_controle': ('helper-controle', 'Controle'),
                        'helper_kernregel': ('helper-kernregel', 'Kernregel'),
                        'helper_tip': ('helper-tip', 'Tip'),
                    }
                    css_cls, label = css_map.get(htype, ('helper-hint', 'Hint'))

                    if htype in ('helper_letop', 'helper_controle', 'helper_kernregel', 'helper_tip'):
                        content = render_callout_text(htext)
                    else:
                        content = render_helper_content(htext)

                    helper_items += f'''            <details class="{css_cls}">
              <summary>{esc(label)}</summary>
              <div class="detail-content">{content}</div>
            </details>
'''
                helpers = f'          <div class="helper-section">\n{helper_items}          </div>\n'

            # Answer toggle
            answer_html = ''
            if q['answer'] or q['uitleg']:
                ans_content = ''
                if q['answer']:
                    ans_lines = q['answer'].replace('\n', '<br>')
                    ans_content += f'              <div class="answer-box">{ans_lines}</div>\n'
                if q['uitleg']:
                    uitleg_text = render_helper_content(q['uitleg'])
                    ans_content += f'              <div class="uitleg-box"><strong>Uitleg:</strong> {uitleg_text}</div>\n'
                answer_html = f'''          <details class="answer-toggle">
            <summary>Toon antwoord</summary>
            <div class="answer-content">
{ans_content}            </div>
          </details>
'''

            questions_html += f'''
        <div class="question-card" id="{q['id']}">
          <div class="question-header">
            <div class="question-label">{esc(q['label'])}</div>
            <div class="question-text">{esc(q['text'])}</div>
          </div>
{fill_in}{helpers}{answer_html}        </div>
'''

        opgave_sections += f'''
      <section class="opgave-section" id="opgave-{opg['nr']}">
        <div class="opgave-header">
          <span class="opgave-num">{opg['nr']}</span>
          <div class="opgave-title-group">
            <div class="opgave-title">{esc(opg['title'])}</div>
            <span class="opgave-badge">{esc(opg.get('domain', 'Marktanalyse'))}</span>
          </div>
        </div>
{ctx}{intro}{questions_html}
      </section>
'''

    # Samenvatting
    samenvatting_html = ''
    if samenvatting:
        rows = ''
        for label, value in samenvatting:
            rows += f'            <tr><td>{esc(label)}</td><td>{esc(value)}</td></tr>\n'
        samenvatting_html = f'''
      <div class="samenvatting-section">
        <div class="samenvatting-caption">&#9632; Samenvattend schema</div>
        <table class="samenvatting-table">
          <tbody>
{rows}          </tbody>
        </table>
      </div>
'''

    n_cards = len(opgaven)
    grid = f'repeat({min(n_cards, 5)}, 1fr)'

    return f'''<!DOCTYPE html>
<html lang="nl">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{para_number} {esc(para_name)} – Begeleide inoefening</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: Arial, sans-serif; background: #F7F8FA; color: #2D3748; line-height: 1.6; min-height: 100vh; }}
  a {{ color: inherit; text-decoration: none; }}
  :root {{
    --navy: #1E2761; --dark: #2D3748; --gray: #718096; --light-gray: #F7F8FA;
    --border-gray: #CBD5E0; --white: #FFFFFF; --sidebar-w: 270px;
    --accent: {dprimary}; --accent-lt: {dlight}; --accent-dk: {ddark};
  }}
  .page-layout {{ display: flex; min-height: 100vh; }}
  .sidebar {{
    width: var(--sidebar-w); flex-shrink: 0; background: var(--white);
    border-right: 1px solid var(--border-gray); position: sticky; top: 0;
    height: 100vh; overflow-y: auto; padding: 1.2rem 0; z-index: 10;
  }}
  .sidebar-header {{ padding: 0.6rem 1.2rem 1rem; border-bottom: 1px solid var(--border-gray); margin-bottom: 0.4rem; }}
  .sidebar-header h2 {{ font-size: 0.82rem; font-weight: bold; color: var(--navy); line-height: 1.3; }}
  .sidebar-header p {{ font-size: 0.7rem; color: var(--gray); margin-top: 0.2rem; }}
  .nav-opgave {{ margin-bottom: 0.15rem; }}
  .nav-opgave-title {{
    display: flex; align-items: center; gap: 0.45rem; padding: 0.45rem 1rem;
    font-size: 0.76rem; font-weight: bold; color: var(--dark); cursor: pointer;
    border-left: 3px solid transparent; transition: background 0.15s;
  }}
  .nav-opgave-title:hover {{ background: #F7FAFC; }}
  .nav-opgave.expanded .nav-opgave-title {{ border-left-color: var(--accent); background: var(--light-gray); }}
  .nav-dot {{ width: 8px; height: 8px; border-radius: 50%; flex-shrink: 0; background: var(--accent); }}
  .nav-opgave-label {{ flex: 1; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }}
  .nav-arrow {{
    width: 14px; height: 14px; flex-shrink: 0; stroke: var(--gray); fill: none;
    stroke-width: 2; stroke-linecap: round; stroke-linejoin: round; transition: transform 0.2s;
  }}
  .nav-opgave.expanded .nav-arrow {{ transform: rotate(180deg); }}
  .nav-questions {{ display: none; padding: 0.1rem 0 0.3rem; }}
  .nav-opgave.expanded .nav-questions {{ display: block; }}
  .nav-q {{
    display: block; padding: 0.3rem 1rem 0.3rem 2.2rem; font-size: 0.72rem;
    color: var(--gray); border-left: 3px solid transparent; overflow: hidden;
    text-overflow: ellipsis; white-space: nowrap; cursor: pointer; transition: background 0.1s;
  }}
  .nav-q:hover {{ color: var(--dark); background: #F7FAFC; }}
  .nav-q.active {{ color: var(--dark); font-weight: bold; border-left-color: var(--accent); background: var(--light-gray); }}
  .sidebar-toggle {{
    display: none; position: fixed; top: 0.75rem; left: 0.75rem; z-index: 20;
    width: 36px; height: 36px; border-radius: 6px; border: 1px solid var(--border-gray);
    background: var(--white); cursor: pointer; align-items: center; justify-content: center;
    box-shadow: 0 1px 4px rgba(0,0,0,0.08);
  }}
  .sidebar-toggle svg {{ width: 20px; height: 20px; stroke: var(--dark); stroke-width: 2; fill: none; }}
  .sidebar-overlay {{ display: none; position: fixed; inset: 0; background: rgba(0,0,0,0.3); z-index: 9; }}
  .content {{ flex: 1; min-width: 0; }}
  .hero {{
    background: var(--navy); color: var(--white); padding: 2.5rem 2rem 1.8rem; position: relative;
  }}
  .hero::before {{ content: ""; position: absolute; top: 0; left: 0; right: 0; height: 5px; background: var(--accent); }}
  .hero-inner {{ max-width: 860px; margin: 0 auto; }}
  .hero-badge {{
    display: inline-block; background: rgba(255,255,255,0.12); border: 1px solid rgba(255,255,255,0.2);
    padding: 0.25rem 0.85rem; border-radius: 4px; font-size: 0.85rem; color: rgba(255,255,255,0.85);
  }}
  .hero h1 {{ font-size: 2.1rem; font-weight: bold; margin: 0.6rem 0 0.3rem; }}
  .hero-sub {{ font-size: 0.95rem; color: rgba(255,255,255,0.65); }}
  .hero-cards {{ display: grid; grid-template-columns: {grid}; gap: 0.7rem; margin-top: 1.5rem; }}
  .hero-card {{
    background: rgba(255,255,255,0.08); border: 1px solid rgba(255,255,255,0.12);
    border-radius: 8px; padding: 0.8rem; cursor: pointer; transition: background 0.15s, transform 0.15s;
    border-top: 3px solid var(--accent);
  }}
  .hero-card:hover {{ background: rgba(255,255,255,0.14); transform: translateY(-2px); }}
  .hero-card-num {{ font-size: 0.7rem; font-weight: bold; opacity: 0.6; margin-bottom: 0.2rem; }}
  .hero-card-title {{ font-size: 0.78rem; font-weight: bold; line-height: 1.3; }}
  .hero-card-count {{ font-size: 0.6rem; opacity: 0.5; margin-top: 0.3rem; }}
  main {{ max-width: 860px; margin: 0 auto; padding: 1.5rem 2rem 3rem; }}
  .opgave-section {{ margin-bottom: 2.5rem; scroll-margin-top: 1rem; }}
  .opgave-header {{
    display: flex; align-items: center; gap: 0.7rem; padding-bottom: 0.7rem;
    margin-bottom: 0.8rem; border-bottom: 3px solid var(--accent);
  }}
  .opgave-num {{
    display: inline-flex; align-items: center; justify-content: center;
    width: 34px; height: 34px; border-radius: 50%; font-size: 0.95rem;
    font-weight: bold; color: var(--white); background: var(--accent); flex-shrink: 0;
  }}
  .opgave-title-group {{ flex: 1; }}
  .opgave-title {{ font-size: 1.2rem; font-weight: bold; color: var(--dark); }}
  .opgave-badge {{
    display: inline-block; font-size: 0.65rem; font-weight: bold; padding: 0.15rem 0.55rem;
    border-radius: 3px; margin-top: 0.15rem; text-transform: uppercase; letter-spacing: 0.03em;
    background: var(--accent-lt); color: var(--accent);
  }}
  .context-box {{
    background: #F8F9FA; border: 1px solid #E2E8F0; border-radius: 6px;
    padding: 0.9rem 1.1rem; margin-bottom: 1.2rem; font-family: Consolas, 'Courier New', monospace;
    font-size: 0.85rem; line-height: 1.7; color: var(--dark); border-left: 4px solid var(--accent);
  }}
  .opgave-intro {{
    font-size: 0.92rem; color: var(--dark); margin-bottom: 1rem; padding: 0.7rem 0.9rem;
    background: #FAFBFC; border-radius: 6px; border-left: 3px solid var(--border-gray);
  }}
  .question-card {{
    background: var(--white); border: 1px solid var(--border-gray); border-radius: 8px;
    margin-bottom: 1rem; overflow: hidden; scroll-margin-top: 1rem;
  }}
  .question-header {{ padding: 0.9rem 1.2rem; border-bottom: 1px solid #EDF2F7; }}
  .question-label {{
    font-size: 0.72rem; font-weight: bold; text-transform: uppercase;
    letter-spacing: 0.05em; color: var(--accent); margin-bottom: 0.3rem;
  }}
  .question-text {{ font-size: 0.92rem; color: var(--dark); line-height: 1.6; }}
  .fill-in {{
    background: #FAFBFC; border: 1px solid #E2E8F0; border-radius: 6px;
    padding: 0.8rem 1rem; margin: 0.6rem 1.2rem; font-family: Consolas, 'Courier New', monospace;
    font-size: 0.85rem; line-height: 1.8; color: var(--dark);
  }}
  .helper-section {{ padding: 0 1.2rem 0.8rem; }}
  details {{ border-radius: 5px; margin-bottom: 0.5rem; overflow: hidden; }}
  details summary {{
    padding: 0.55rem 0.8rem; font-size: 0.82rem; font-weight: 600; cursor: pointer;
    user-select: none; list-style: none; display: flex; align-items: center; gap: 0.4rem;
    border-radius: 5px; transition: background 0.15s;
  }}
  details summary::-webkit-details-marker {{ display: none; }}
  details summary::before {{ content: "▸"; font-size: 0.7rem; transition: transform 0.2s; flex-shrink: 0; }}
  details[open] summary::before {{ transform: rotate(90deg); }}
  details .detail-content {{ padding: 0.5rem 0.8rem 0.7rem; font-size: 0.85rem; line-height: 1.7; }}
  .helper-denkstappen summary {{ background: #E8F4FD; color: #1565C0; }}
  .helper-denkstappen summary:hover {{ background: #D0EAFB; }}
  .helper-denkstappen .detail-content {{ background: #F1F8FE; color: #0D47A1; }}
  .helper-hint summary {{ background: #FFF8E1; color: #F57F17; }}
  .helper-hint summary:hover {{ background: #FFF0C0; }}
  .helper-hint .detail-content {{ background: #FFFCF0; color: #E65100; }}
  .helper-formule summary {{ background: #F3E5F5; color: #6A1B9A; }}
  .helper-formule summary:hover {{ background: #E8D5EB; }}
  .helper-formule .detail-content {{ background: #F9F0FB; color: #4A148C; font-family: Consolas, 'Courier New', monospace; }}
  .helper-letop summary {{ background: #FFF3E0; color: #BF360C; }}
  .helper-letop summary:hover {{ background: #FFE8CC; }}
  .helper-letop .detail-content {{ background: #FFF8F0; color: #BF360C; }}
  .helper-controle summary {{ background: #E8F5E9; color: #1B5E20; }}
  .helper-controle summary:hover {{ background: #D5ECD7; }}
  .helper-controle .detail-content {{ background: #F0FAF0; color: #1B5E20; }}
  .helper-kernregel summary {{ background: #FFF8E1; color: #5D4037; }}
  .helper-kernregel summary:hover {{ background: #FFF0C0; }}
  .helper-kernregel .detail-content {{ background: #FFFCF0; color: #5D4037; }}
  .helper-tip summary {{ background: #EBF5FB; color: #154360; }}
  .helper-tip summary:hover {{ background: #D6EAF8; }}
  .helper-tip .detail-content {{ background: #F4FAFE; color: #154360; }}
  .answer-toggle {{ border-top: 1px solid #EDF2F7; }}
  .answer-toggle summary {{
    padding: 0.7rem 1.2rem; font-size: 0.85rem; font-weight: bold; color: var(--accent);
    cursor: pointer; user-select: none; list-style: none; display: flex; align-items: center;
    gap: 0.4rem; background: var(--accent-lt); transition: background 0.15s;
  }}
  .answer-toggle summary::-webkit-details-marker {{ display: none; }}
  .answer-toggle summary::before {{ content: "▸"; font-size: 0.75rem; transition: transform 0.2s; }}
  .answer-toggle[open] summary::before {{ transform: rotate(90deg); }}
  .answer-toggle[open] summary {{ background: #D4E6F1; }}
  .answer-content {{ padding: 1rem 1.2rem; border-top: 1px solid #D6EAF8; }}
  .answer-box {{
    background: #F0FFF0; border: 1px solid #C8E6C9; border-radius: 6px;
    padding: 0.8rem 1rem; margin-bottom: 0.6rem; font-family: Consolas, 'Courier New', monospace;
    font-size: 0.85rem; line-height: 1.8; color: #1B5E20;
  }}
  .uitleg-box {{
    background: #FFFDE7; border-left: 3px solid #AFB42B; border-radius: 0 5px 5px 0;
    padding: 0.7rem 0.9rem; font-size: 0.85rem; line-height: 1.6; color: #827717;
  }}
  .samenvatting-section {{
    background: var(--white); border: 1px solid var(--border-gray);
    border-radius: 8px; overflow: hidden; margin-top: 2rem;
  }}
  .samenvatting-caption {{
    background: var(--accent); color: var(--white); padding: 0.7rem 1rem;
    font-weight: bold; font-size: 0.9rem;
  }}
  .samenvatting-table {{ width: 100%; border-collapse: collapse; font-size: 0.88rem; }}
  .samenvatting-table td {{ padding: 0.55rem 1rem; border-bottom: 1px solid #EDF2F7; vertical-align: top; }}
  .samenvatting-table tr:last-child td {{ border-bottom: none; }}
  .samenvatting-table td:first-child {{ font-weight: 600; color: var(--dark); width: 28%; }}
  .samenvatting-table td:last-child {{ color: #4A5568; }}
  @media (max-width: 1100px) {{ .hero-cards {{ grid-template-columns: repeat(3, 1fr); }} }}
  @media (max-width: 768px) {{
    .sidebar {{
      position: fixed; left: -290px; top: 0; height: 100vh; width: 290px;
      transition: left 0.25s ease; box-shadow: 2px 0 8px rgba(0,0,0,0.1);
    }}
    .sidebar.open {{ left: 0; }}
    .sidebar-toggle {{ display: flex; }}
    .sidebar-overlay.show {{ display: block; }}
    .hero {{ padding: 2rem 1.2rem 1.5rem; }}
    .hero h1 {{ font-size: 1.6rem; }}
    .hero-cards {{ grid-template-columns: 1fr 1fr; gap: 0.6rem; }}
    main {{ padding: 1rem 1.2rem 2rem; }}
  }}
  @media (max-width: 480px) {{ .hero-cards {{ grid-template-columns: 1fr; }} }}
</style>
</head>
<body>
<button class="sidebar-toggle" id="sidebarToggle" aria-label="Menu openen">
  <svg viewBox="0 0 24 24"><line x1="3" y1="6" x2="21" y2="6"/><line x1="3" y1="12" x2="21" y2="12"/><line x1="3" y1="18" x2="21" y2="18"/></svg>
</button>
<div class="sidebar-overlay" id="sidebarOverlay"></div>
<div class="page-layout">
  <nav class="sidebar" id="sidebar">
    <div class="sidebar-header">
      <h2>{para_number} {esc(para_name)}</h2>
      <p>Begeleide inoefening</p>
    </div>
{sidebar_items}
  </nav>
  <div class="content">
    <header class="hero">
      <div class="hero-inner">
        <span class="hero-badge">{para_number} &middot; Begeleide inoefening</span>
        <h1>{esc(para_name)} &mdash; Begeleide inoefening</h1>
        <p class="hero-sub">Oefeningen met denkstappen, formule-herinneringen en hints.</p>
        <div class="hero-cards">
{hero_cards}
        </div>
      </div>
    </header>
    <main>
{opgave_sections}
{samenvatting_html}
    </main>
  </div>
</div>
<script>
(function() {{
  var sidebar = document.getElementById('sidebar');
  var overlay = document.getElementById('sidebarOverlay');
  var toggle = document.getElementById('sidebarToggle');
  document.querySelectorAll('[data-toggle="opgave"]').forEach(function(title) {{
    title.addEventListener('click', function() {{
      var opgave = this.closest('.nav-opgave');
      opgave.classList.toggle('expanded');
      var targetId = this.dataset.scroll;
      if (targetId) {{
        var target = document.getElementById(targetId);
        if (target) {{
          target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
          sidebar.classList.remove('open');
          overlay.classList.remove('show');
        }}
      }}
    }});
  }});
  document.querySelectorAll('.nav-q[data-q]').forEach(function(el) {{
    el.addEventListener('click', function(e) {{
      e.preventDefault();
      var target = document.getElementById(el.dataset.q);
      if (target) {{
        target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
        sidebar.classList.remove('open');
        overlay.classList.remove('show');
      }}
    }});
  }});
  document.querySelectorAll('[data-target]').forEach(function(el) {{
    el.addEventListener('click', function(e) {{
      var target = document.getElementById(el.dataset.target);
      if (target) {{
        e.preventDefault();
        target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
      }}
    }});
  }});
  var navQs = document.querySelectorAll('.nav-q[data-q]');
  var questionCards = document.querySelectorAll('.question-card');
  var observer = new IntersectionObserver(function(entries) {{
    entries.forEach(function(entry) {{
      if (entry.isIntersecting) {{
        navQs.forEach(function(n) {{ n.classList.remove('active'); }});
        var active = document.querySelector('.nav-q[data-q="' + entry.target.id + '"]');
        if (active) {{
          active.classList.add('active');
          var parent = active.closest('.nav-opgave');
          if (parent && !parent.classList.contains('expanded')) parent.classList.add('expanded');
        }}
      }}
    }});
  }}, {{ rootMargin: '-10% 0px -80% 0px' }});
  questionCards.forEach(function(q) {{ observer.observe(q); }});
  toggle.addEventListener('click', function() {{
    sidebar.classList.toggle('open');
    overlay.classList.toggle('show');
  }});
  overlay.addEventListener('click', function() {{
    sidebar.classList.remove('open');
    overlay.classList.remove('show');
  }});
}})();
</script>
</body>
</html>'''


def find_paragraph_info(folder_path):
    basename = os.path.basename(folder_path.rstrip('/\\'))
    parts = basename.split(' ', 1)
    number = parts[0]
    if ' - ' in basename:
        name = basename.split(' - ', 1)[1]
    else:
        name = parts[1] if len(parts) > 1 else basename
    return number, name


def process_paragraph(para_folder):
    para_number, para_name = find_paragraph_info(para_folder)

    bi_folder = os.path.join(para_folder, '3. Oefenen', 'begeleide inoefening')
    if not os.path.isdir(bi_folder):
        print(f'  SKIP {para_number}: no begeleide inoefening folder')
        return False

    vragen_files = glob.glob(os.path.join(bi_folder, '*vragen*.docx'))
    antwoorden_files = glob.glob(os.path.join(bi_folder, '*antwoorden*.docx'))

    if not vragen_files:
        print(f'  SKIP {para_number}: no vragen.docx')
        return False

    vragen_stream = parse_bi_document(vragen_files[0], is_answers=False)
    antwoorden_stream = parse_bi_document(antwoorden_files[0], is_answers=True) if antwoorden_files else None

    if vragen_stream is None:
        print(f'  ERROR {para_number}: could not parse vragen')
        return False

    opgaven, samenvatting = merge_vragen_antwoorden(vragen_stream, antwoorden_stream)

    if not opgaven:
        print(f'  SKIP {para_number}: no opgaven found')
        return False

    total_q = sum(len(o['questions']) for o in opgaven)
    total_ans = sum(1 for o in opgaven for q in o['questions'] if q['answer'])

    html_content = generate_html(opgaven, samenvatting, para_number, para_name)

    html_filename = f'{para_number} {para_name} \u2013 begeleide inoefening.html'
    html_path = os.path.join(bi_folder, html_filename)

    with open(long_path(html_path), 'w', encoding='utf-8') as f:
        f.write(html_content)

    print(f'  OK {para_number} {para_name} ({len(opgaven)} opgaven, {total_q} vragen, {total_ans} antwoorden)')
    return True


def main():
    if len(sys.argv) < 2:
        print('Usage: python convert_begeleide_inoefening.py [path] or --all')
        sys.exit(1)

    if sys.argv[1] == '--all':
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        folders = sorted(glob.glob(os.path.join(base, '3.*/3.*.*')))
        print(f'Found {len(folders)} paragraph folders')
        success = 0
        for folder in folders:
            if os.path.isdir(folder):
                if process_paragraph(folder):
                    success += 1
        print(f'\nDone: {success}/{len(folders)} converted')
    else:
        process_paragraph(sys.argv[1])


if __name__ == '__main__':
    main()
